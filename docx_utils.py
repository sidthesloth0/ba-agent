import streamlit as st
import re
from docx import Document
from docx.shared import Inches
import io
import base64
import requests
import google.generativeai as genai
from prompts import MERMAID_PROMPT, TRD_PROMPT



def generate_mermaid_req_doc(md_text):
    """Generates a technical requirements document in Mermaid.js format."""
    model = genai.GenerativeModel('gemini-1.5-flash')

    prompt = MERMAID_PROMPT.format(md_text=md_text)

    try:
        response = model.generate_content(prompt)
        # Clean the response to ensure it's valid Mermaid code
        mermaid_code = response.text
        # The model might still wrap the code in markdown fences. Let's extract it.
        match = re.search(r"```mermaid\n(.*?)\n```", mermaid_code, re.DOTALL)
        if match:
            mermaid_code = match.group(1).strip()
        else:
            mermaid_code = mermaid_code.strip()

        # Fallback to remove parentheses only from within node labels
        def clean_label(match):
            label_content = match.group(1)
            cleaned_content = re.sub(r'[\(\)]', '', label_content)
            return f'[{cleaned_content}]'

        mermaid_code = re.sub(r'\[(.*?)\]', clean_label, mermaid_code)

        usage = response.usage_metadata
        token_info = {
            "prompt": usage.prompt_token_count,
            "output": usage.candidates_token_count,
            "total": usage.total_token_count
        }
        return mermaid_code.strip(), token_info

    except Exception as e:
        st.error(f"An error occurred during Mermaid diagram generation: {e}")
        return None, None


def generate_trd_content(md_text):
    """Generates a full technical requirements document based on the business plan."""
    model = genai.GenerativeModel('gemini-1.5-flash')

    prompt = TRD_PROMPT.format(md_text=md_text)

    try:
        response = model.generate_content(prompt)
        usage = response.usage_metadata
        token_info = {
            "prompt": usage.prompt_token_count,
            "output": usage.candidates_token_count,
            "total": usage.total_token_count
        }
        return response.text, token_info
    except Exception as e:
        st.error(f"An error occurred during TRD content generation: {e}")
        return None, None

def add_md_to_doc(document, markdown_text):
    """
    Parses markdown text and adds it to a python-docx document.
    Handles headings, lists, and tables.
    """
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if re.match(r'^#\s', line):
            document.add_heading(line.lstrip('# ').strip(), level=1)
        elif re.match(r'^##\s', line):
            document.add_heading(line.lstrip('## ').strip(), level=2)
        elif re.match(r'^###\s', line):
            document.add_heading(line.lstrip('### ').strip(), level=3)
        elif line.startswith(('* ', '- ')):
            document.add_paragraph(line.lstrip('* -'), style='List Bullet')
        elif line.startswith('|') and line.endswith('|'):
            # Check if the next line is a table separator, indicating a table
            if i + 1 < len(lines) and re.match(r'^\s*\|[:\- |]+\|\s*$', lines[i+1].strip()):
                header = [h.strip() for h in line.strip('|').split('|')]
                num_cols = len(header)
                
                table = document.add_table(rows=1, cols=num_cols)
                table.style = 'Table Grid'
                
                # Populate header row
                hdr_cells = table.rows[0].cells
                for j, col_header in enumerate(header):
                    hdr_cells[j].text = col_header
                
                # Skip header and separator lines
                i += 2
                
                # Process data rows
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_data = [cell.strip() for cell in lines[i].strip('|').split('|')]
                    row_cells = table.add_row().cells
                    for j in range(min(num_cols, len(row_data))):
                        row_cells[j].text = row_data[j]
                    i += 1
                continue # Continue to next line after table processing
            else:
                document.add_paragraph(line)
        else:
            document.add_paragraph(line)
        
        i += 1

def create_trd_word_document(trd_content, mermaid_code, epics_user_stories=None):
    """Creates a Word document with TRD content, a Mermaid diagram, and epics/user stories."""
    try:
        document = Document()
        
        # Add the TRD content first
        add_md_to_doc(document, trd_content)

        # Add the diagram
        document.add_page_break()
        document.add_heading('System Architecture Diagram', level=1)
        
        # Generate the image URL from mermaid.ink
        graphbytes = mermaid_code.encode("utf8")
        base64_bytes = base64.b64encode(graphbytes)
        base64_string = base64_bytes.decode("utf-8")
        img_url = f"https://mermaid.ink/img/{base64_string}"

        # Fetch the image
        response = requests.get(img_url)
        if response.status_code != 200:
            st.error(f"Failed to fetch diagram image from mermaid.ink. Status: {response.status_code}")
            return None

        image_stream = io.BytesIO(response.content)
        document.add_picture(image_stream, width=Inches(6.0))

        # Add Epics and User Stories if they exist
        if epics_user_stories:
            document.add_page_break()
            document.add_heading('Epics and User Stories', level=1)
            add_md_to_doc(document, epics_user_stories)
        
        # Save the document to a byte stream
        doc_stream = io.BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream

    except Exception as e:
        st.error(f"An error occurred while creating the Word document: {e}")
        return None