import streamlit as st
from dotenv import load_dotenv
import google.generativeai as genai
import os
import pymupdf4llm
import fitz  # PyMuPDF
from PIL import Image
import io
import re
from streamlit_mermaid import st_mermaid
import base64
from docx import Document
from docx.shared import Inches
import requests

load_dotenv(override=True)

st.title("PDF Business Plan Analyzer")

# Custom CSS to reduce vertical spacing for a more compact layout
st.markdown("""
<style>
    /* Main container adjustments */
    .main .block-container {
        padding-top: 2rem; /* A bit of space at the top */
    }

    /* Heading adjustments for compactness */
    h1 {
        margin-bottom: 0.5rem !important; /* Space after main title */
    }
    h2, h3 { /* Affects st.subheader and st.header */
        margin-top: 2rem !important; /* Space before section headers */
    }

    /* Sidebar adjustments for compactness */
    [data-testid="stSidebar"] div.stMarkdown {
        padding-top: 0px !important;
        padding-bottom: 0px !important;
    }
    [data-testid="stSidebar"] h2 {
        margin-top: 1rem; /* Space above section headers */
        margin-bottom: 0.25rem; /* Reduce space after sidebar headers */
    }
    /* Custom style for the token display in the sidebar */
    .sidebar-token-usage {
        font-size: 0.9rem;
        color: grey;
        padding: 0px;
        margin-bottom: 0.5rem; /* Reduced bottom margin */
    }
</style>
""", unsafe_allow_html=True)

gemini_api_key = os.getenv("GOOGLE_API_KEY")
if not gemini_api_key or gemini_api_key.strip() == "":
    st.error("GOOGLE_API_KEY environment variable not found.")
    st.stop()

genai.configure(api_key=gemini_api_key)

def extract_headings(markdown_text):
    """Extracts headings from markdown text for a table of contents."""
    headings = []
    for line in markdown_text.split('\n'):
        match = re.match(r'^(#{1,6})\s+(.+)', line.strip())
        if match:
            level = len(match.group(1))
            title = match.group(2).strip()
            anchor_id = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '-').lower()
            headings.append({'level': level, 'title': title, 'anchor': anchor_id})
    return headings

def create_sidebar_toc(headings):
    # creates a sidebar table of contents from extracted headings
    if headings:
        st.sidebar.markdown("## Table of Contents")
        for heading in headings:
            indent = "  " * (heading['level'] - 1)
            st.sidebar.markdown(f"{indent}- [{heading['title']}](#{heading['anchor']})")

def analyze_with_gemini(md_text, images):
    """Analyzes the business plan text and images using Gemini."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    prompt = f"""
    As a business analyst, analyze the provided business plan.
    Provide a concise analysis covering:
    1. Summary
    2. Strengths
    3. Weaknesses
    4. Actionable suggestions
    5. Image analysis

    Business Plan Content:
    **Text (Markdown):**
    {md_text}
    """
    
    content_parts = [prompt]
    
    if images:
        content_parts.append("\n**Images:**")
        for img in images:
            content_parts.append(img)
    
    try:
        response = model.generate_content(content_parts)
        usage = response.usage_metadata
        token_info = {
            "prompt": usage.prompt_token_count,
            "output": usage.candidates_token_count,
            "total": usage.total_token_count
        }
        return response.text, token_info
    except Exception as e:
        st.error(f"An error occurred during Gemini analysis: {e}")
        return None, None

def generate_mermaid_req_doc(md_text):
    """Generates a technical requirements document in Mermaid.js format."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    prompt = f"""
    As a technical architect, create a Mermaid.js diagram from the business plan.
    - Diagram must be top-down (`graph TD`).
    - Outline system architecture, components, and user flow.
    - Output ONLY raw Mermaid.js code using standard Mermaid syntax like `id[Description]`.
    - CRITICAL: The text inside the node brackets (e.g., "Description") MUST NOT contain any parentheses or brackets. For example, instead of `A[Node with (parentheses)]`, write `A[Node with parentheses]`.
    - No explanations, markdown fences, custom styles, or link labels.

    Business Plan:
    {md_text}
    """
    
    try:
        response = model.generate_content(prompt)
        # Clean the response to ensure it's valid Mermaid code
        mermaid_code = response.text
        # The model might still wrap the code in markdown fences. Let's extract it.
        match = re.search(r"```mermaid\n(.*?)\n```", mermaid_code, re.DOTALL)
        if match:
            mermaid_code = match.group(1).strip()
        else:
            mermaid_code = mermaid_code.strip()

        # Fallback to remove parentheses only from within node labels
        def clean_label(match):
            label_content = match.group(1)
            cleaned_content = re.sub(r'[\(\)]', '', label_content)
            return f'[{cleaned_content}]'

        mermaid_code = re.sub(r'\[(.*?)\]', clean_label, mermaid_code)

        usage = response.usage_metadata
        token_info = {
            "prompt": usage.prompt_token_count,
            "output": usage.candidates_token_count,
            "total": usage.total_token_count
        }
        return mermaid_code.strip(), token_info

    except Exception as e:
        st.error(f"An error occurred during Mermaid diagram generation: {e}")
        return None, None

def generate_trd_content(md_text):
    """Generates a full technical requirements document based on the business plan."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    prompt = f"""
    As a senior business analyst, create a comprehensive Technical Requirements Document (TRD) in Markdown format based on the provided business plan.

    **TRD Structure:**
    Follow this structure precisely. ONLY include sections if the business plan provides relevant information.

    # 1. Introduction
    ## Project Overview
    (A brief summary of the project's purpose and scope)
    ## Objectives
    (List the key goals of the project)

    # 2. System Overview
    ## High-Level Architecture
    (Describe the proposed architecture)
    ## System Components
    (Detail the main components or modules)

    # 3. Functional Requirements
    (Create a plain text table with these columns: | FR-ID | Requirement | Description |)

    # 4. UI Screen Specifications
    (ONLY include this section if UI/UX details are mentioned. Use '## Screen: [Screen Name]' as a sub-header for each screen.)

    # 5. User Roles and Permissions
    (ONLY include this section if user roles are described. Create a plain text table: | Role | Permissions |)

    # 6. Technical Specifications
    (ONLY include this section if technical details are provided.)
    ## Database Schema
    (Propose a high-level database schema if applicable)
    ## API Endpoints
    (Suggest key API endpoints if applicable)

    # 7. Non-Functional Requirements
    (ONLY include this section if NFRs like performance, security, or scalability are mentioned.)
    ## Performance
    (Describe performance expectations)
    ## Security
    (Outline security considerations)
    ## Scalability
    (Mention scalability requirements)

    **Business Plan Content:**
    {md_text}
    """
    
    try:
        response = model.generate_content(prompt)
        usage = response.usage_metadata
        token_info = {
            "prompt": usage.prompt_token_count,
            "output": usage.candidates_token_count,
            "total": usage.total_token_count
        }
        return response.text, token_info
    except Exception as e:
        st.error(f"An error occurred during TRD content generation: {e}")
        return None, None

def summarize_text(md_text):
    """Generates a concise summary of the text."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    prompt = f"""
    Please provide a concise summary of the following business plan text.
    Focus on the key points, objectives, and strategies.
    The summary should be a few paragraphs long.

    Business Plan Text:
    {md_text}
    """
    
    try:
        response = model.generate_content(prompt)
        usage = response.usage_metadata
        token_info = {
            "prompt": usage.prompt_token_count,
            "output": usage.candidates_token_count,
            "total": usage.total_token_count
        }
        return response.text, token_info
    except Exception as e:
        st.error(f"An error occurred during summarization: {e}")
        return None, None

def add_md_to_doc(document, markdown_text):
    """
    Parses markdown text and adds it to a python-docx document.
    Handles headings, lists, and tables.
    """
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if re.match(r'^#\s', line):
            document.add_heading(line.lstrip('# ').strip(), level=1)
        elif re.match(r'^##\s', line):
            document.add_heading(line.lstrip('## ').strip(), level=2)
        elif re.match(r'^###\s', line):
            document.add_heading(line.lstrip('### ').strip(), level=3)
        elif line.startswith(('* ', '- ')):
            document.add_paragraph(line.lstrip('* -'), style='List Bullet')
        elif line.startswith('|') and line.endswith('|'):
            # Check if the next line is a table separator, indicating a table
            if i + 1 < len(lines) and re.match(r'^\s*\|[:\- |]+\|\s*$', lines[i+1].strip()):
                header = [h.strip() for h in line.strip('|').split('|')]
                num_cols = len(header)
                
                table = document.add_table(rows=1, cols=num_cols)
                table.style = 'Table Grid'
                
                # Populate header row
                hdr_cells = table.rows[0].cells
                for j, col_header in enumerate(header):
                    hdr_cells[j].text = col_header
                
                # Skip header and separator lines
                i += 2
                
                # Process data rows
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_data = [cell.strip() for cell in lines[i].strip('|').split('|')]
                    row_cells = table.add_row().cells
                    for j in range(min(num_cols, len(row_data))):
                        row_cells[j].text = row_data[j]
                    i += 1
                continue # Continue to next line after table processing
            else:
                document.add_paragraph(line)
        else:
            document.add_paragraph(line)
        
        i += 1

def create_trd_word_document(trd_content, mermaid_code):
    """Creates a Word document with TRD content and a Mermaid diagram."""
    try:
        document = Document()
        
        # Add the TRD content first
        add_md_to_doc(document, trd_content)

        # Add the diagram
        document.add_page_break()
        document.add_heading('System Architecture Diagram', level=1)
        
        # Generate the image URL from mermaid.ink
        graphbytes = mermaid_code.encode("utf8")
        base64_bytes = base64.b64encode(graphbytes)
        base64_string = base64_bytes.decode("utf-8")
        img_url = f"https://mermaid.ink/img/{base64_string}"

        # Fetch the image
        response = requests.get(img_url)
        if response.status_code != 200:
            st.error(f"Failed to fetch diagram image from mermaid.ink. Status: {response.status_code}")
            return None

        image_stream = io.BytesIO(response.content)
        document.add_picture(image_stream, width=Inches(6.0))
        
        # Save the document to a byte stream
        doc_stream = io.BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream

    except Exception as e:
        st.error(f"An error occurred while creating the Word document: {e}")
        return None

# --- App Logic with Session State ---

# Initialize session state to store processed data
if "md_text" not in st.session_state:
    st.session_state.md_text = None
    st.session_state.image_list = []
    st.session_state.last_filename = None
    st.session_state.token_counts = {"prompt": 0, "output": 0, "total": 0}
    st.session_state.analysis = None
    st.session_state.mermaid_code = None
    st.session_state.summary = None
    st.session_state.trd_content = None

uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")

# 1. Process the file only if it's a new upload
if uploaded_file is not None and uploaded_file.name != st.session_state.last_filename:
    # Reset token counts for the new file session
    st.session_state.token_counts = {"prompt": 0, "output": 0, "total": 0}
    st.session_state.analysis = None
    st.session_state.mermaid_code = None
    st.session_state.summary = None
    st.session_state.trd_content = None
    st.session_state.last_filename = uploaded_file.name
    temp_pdf_path = "temp_uploaded.pdf"
    try:
        with st.spinner("Processing PDF..."):
            # Write to a temporary file to be processed
            with open(temp_pdf_path, "wb") as f:
                f.write(uploaded_file.getvalue())

            # Extract markdown text and store in session state
            st.session_state.md_text = pymupdf4llm.to_markdown(temp_pdf_path)

            # Extract images and store in session state
            image_list = []
            with fitz.open(temp_pdf_path) as doc:
                for page in doc:
                    images = page.get_images(full=True)
                    for img_index, img in enumerate(images):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image = Image.open(io.BytesIO(image_bytes))
                        image_list.append(image)
            st.session_state.image_list = image_list

    except Exception as e:
        st.error("An error occurred while processing the PDF. See details below.")
        st.exception(e)
        # Reset state in case of an error
        st.session_state.md_text = None
        st.session_state.image_list = []
        st.session_state.last_filename = None
        st.session_state.analysis = None
        st.session_state.mermaid_code = None
        st.session_state.summary = None
        st.session_state.trd_content = None
    finally:
        # Clean up the temporary file immediately after processing
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)

# 2. Display content and analysis button if a file has been successfully processed
if st.session_state.md_text:
    # --- Sidebar ---
    st.sidebar.markdown("## Token Usage")
    st.sidebar.markdown(
        f"""
        <div class="sidebar-token-usage">
            <strong>Total:</strong> {st.session_state.token_counts['total']}<br>
            <strong>Prompt:</strong> {st.session_state.token_counts['prompt']}<br>
            <strong>Output:</strong> {st.session_state.token_counts['output']}
        </div>
        """,
        unsafe_allow_html=True
    )

    st.sidebar.markdown("## Navigation")
    st.sidebar.markdown("[Extracted Text](#extracted-markdown-text)")
    st.sidebar.markdown("[Extracted Images](#extracted-images)")
    st.sidebar.markdown("[AI Analysis](#gemini-analysis)")

    headings = extract_headings(st.session_state.md_text)
    create_sidebar_toc(headings)

    # --- Main Content ---
    st.subheader("Extracted Text", anchor="extracted-markdown-text")
    st.markdown(st.session_state.md_text, unsafe_allow_html=True)

    st.subheader("Extracted Images", anchor="extracted-images")
    if st.session_state.image_list:
        for i, image in enumerate(st.session_state.image_list):
            st.image(image, caption=f"Image {i+1}")
    else:
        st.info("No images found in the PDF.")

    # --- Summarization Section ---
    st.subheader("Token Optimization: Summary")
    if st.button("Generate Summary"):
        with st.spinner("Generating summary..."):
            summary, token_info = summarize_text(st.session_state.md_text)
            if summary and token_info:
                st.session_state.summary = summary
                # Update token counts
                st.session_state.token_counts["prompt"] += token_info["prompt"]
                st.session_state.token_counts["output"] += token_info["output"]
                st.session_state.token_counts["total"] += token_info["total"]
                st.rerun()

    if st.session_state.summary:
        st.markdown("### Generated Summary")
        with st.container(height=200):
            st.markdown(st.session_state.summary)

    # --- Analysis and TRD Section ---
    st.subheader("AI Analysis & Technical Document Generation", anchor="gemini-analysis")

    use_summary = False
    if st.session_state.summary:
        use_summary = st.checkbox("Use generated summary for analysis and TRD", value=True)

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Analyze Business Plan"):
            with st.spinner("Analyzing with Gemini..."):
                input_text = st.session_state.summary if use_summary else st.session_state.md_text
                images_to_analyze = [] if use_summary else st.session_state.image_list
                
                analysis, token_info = analyze_with_gemini(input_text, images_to_analyze)
                if analysis and token_info:
                    st.session_state.analysis = analysis
                    # Update token counts
                    st.session_state.token_counts["prompt"] += token_info["prompt"]
                    st.session_state.token_counts["output"] += token_info["output"]
                    st.session_state.token_counts["total"] += token_info["total"]
                    st.rerun()

    with col2:
        if st.button("Generate Technical Requirements Document"):
            with st.spinner("Generating Technical Requirements Document... This may take a moment."):
                input_text = st.session_state.summary if use_summary else st.session_state.md_text
                
                # Generate Mermaid Diagram
                mermaid_code, mermaid_token_info = generate_mermaid_req_doc(input_text)
                
                # Generate TRD Content
                trd_content, trd_token_info = generate_trd_content(input_text)

                if mermaid_code and mermaid_token_info and trd_content and trd_token_info:
                    st.session_state.mermaid_code = mermaid_code
                    st.session_state.trd_content = trd_content
                    
                    # Update token counts by aggregating both calls
                    st.session_state.token_counts["prompt"] += mermaid_token_info["prompt"] + trd_token_info["prompt"]
                    st.session_state.token_counts["output"] += mermaid_token_info["output"] + trd_token_info["output"]
                    st.session_state.token_counts["total"] += mermaid_token_info["total"] + trd_token_info["total"]
                    
                    st.rerun()

    if st.session_state.analysis:
        st.markdown("### Business Analysis")
        st.markdown(st.session_state.analysis)

    if st.session_state.mermaid_code and st.session_state.trd_content:
        st.markdown("### Technical Requirements Document")
        st.success("Technical Requirements Document generated successfully!")
        
        # Display TRD content
        with st.expander("View Generated TRD Content", expanded=True):
            st.markdown(st.session_state.trd_content)

        # Render the diagram in the app
        st.markdown("#### System Architecture Diagram")
        st_mermaid(st.session_state.mermaid_code)

        doc_stream = create_trd_word_document(st.session_state.trd_content, st.session_state.mermaid_code)

        if doc_stream:
            st.download_button(
                label="Download Full TRD (Word Document)",
                data=doc_stream,
                file_name="Technical_Requirements_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with st.expander("View and Copy Mermaid.js Code"):
            st.code(st.session_state.mermaid_code, language="mermaid")
